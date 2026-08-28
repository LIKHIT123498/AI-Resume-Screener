import io
from pypdf import PdfReader
from docx import Document

def extract_text_from_file(file_content: bytes, filename: str) -> str:
    """
    Robustly extracts text from uploaded PDF or DOCX bytes, handling paragraphs, runs, tables, and hyperlinks.
    """
    text = ""
    try:
        if filename.lower().endswith(".pdf"):
            reader = PdfReader(io.BytesIO(file_content))
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
                    
        elif filename.lower().endswith(".docx"):
            doc = Document(io.BytesIO(file_content))
            
            # Extract text from paragraphs and inline runs (captures formatted text & links)
            for paragraph in doc.paragraphs:
                para_text = ""
                for run in paragraph.runs:
                    if run.text:
                        para_text += run.text
                # Fallback to paragraph.text if runs came up empty
                if not para_text.strip():
                    para_text = paragraph.text
                if para_text.strip():
                    text += para_text + "\n"
                    
            # Extract text from tables if any exist
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = "".join([p.text for p in cell.paragraphs])
                        if cell_text.strip():
                            text += cell_text + "\n"
                            
        else:
            raise ValueError("Unsupported file format.")
            
    except Exception as e:
        print(f"Error parsing file {filename}: {e}")
        return ""
    
    return text.strip()